#!/usr/bin/env python3
"""
Convert PROJECT_REPORT.md to PROJECT_REPORT.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read the markdown file
with open('PROJECT_REPORT.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Split content into lines
lines = content.split('\n')

i = 0
while i < len(lines):
    line = lines[i]
    
    # Handle headings
    if line.startswith('# '):
        title = line.replace('# ', '')
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
        i += 1
        continue
    
    elif line.startswith('## '):
        title = line.replace('## ', '')
        heading = doc.add_heading(title, level=2)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True
        i += 1
        continue
    
    elif line.startswith('### '):
        title = line.replace('### ', '')
        heading = doc.add_heading(title, level=3)
        for run in heading.runs:
            run.font.size = Pt(12)
            run.font.bold = True
        i += 1
        continue
    
    elif line.startswith('#### '):
        title = line.replace('#### ', '')
        heading = doc.add_heading(title, level=4)
        for run in heading.runs:
            run.font.size = Pt(11)
            run.font.bold = True
        i += 1
        continue
    
    # Handle code blocks
    elif line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        code_text = '\n'.join(code_lines)
        code_para = doc.add_paragraph(code_text, style='List Number')
        code_para_format = code_para.paragraph_format
        code_para_format.left_indent = Inches(0.5)
        
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        i += 1
        continue
    
    # Handle bullet points and lists
    elif line.strip().startswith('- '):
        bullet_text = line.replace('- ', '').strip()
        para = doc.add_paragraph(bullet_text, style='List Bullet')
        i += 1
        continue
    
    elif line.strip().startswith('◦ '):
        bullet_text = line.replace('◦ ', '').strip()
        para = doc.add_paragraph(bullet_text, style='List Bullet 2')
        i += 1
        continue
    
    elif line.strip().startswith('▪ '):
        bullet_text = line.replace('▪ ', '').strip()
        para = doc.add_paragraph(bullet_text, style='List Bullet')
        i += 1
        continue
    
    # Handle numbered list
    elif re.match(r'^\d+\. ', line.strip()):
        numbered_text = re.sub(r'^\d+\. ', '', line.strip())
        para = doc.add_paragraph(numbered_text, style='List Number')
        i += 1
        continue
    
    # Handle links
    elif line.strip().startswith('[') and '](http' in line:
        # Extract link text and URL
        match = re.search(r'\[([^\]]+)\]\(([^\)]+)\)', line)
        if match:
            link_text = match.group(1)
            link_url = match.group(2)
            para = doc.add_paragraph()
            run = para.add_run(link_text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        i += 1
        continue
    
    # Handle empty lines
    elif line.strip() == '':
        doc.add_paragraph()
        i += 1
        continue
    
    # Handle regular text
    else:
        text = line.strip()
        if text:
            para = doc.add_paragraph(text)
            para.paragraph_format.line_spacing = 1.15
        i += 1
        continue

# Save the document
doc.save('PROJECT_REPORT.docx')
print("✓ Word document created: PROJECT_REPORT.docx")
